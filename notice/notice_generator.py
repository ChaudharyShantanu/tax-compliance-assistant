# notice/notice_generator.py
# Builds the final notice as DOCX and PDF.
# PDF uses reportlab (more widely available than fpdf2).
#
# Dependencies:
#   python-docx  — DOCX generation
#   reportlab    — PDF generation  (pip install reportlab)

from datetime import date
from io import BytesIO

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


# ── Internal helpers ──────────────────────────────────────────────────────

_SKIP_PREFIXES = (
    "sir/madam", "to,", "subject:", "dear sir",
    "income tax department", "government of india",
)

def _body_lines(notice_text: str):
    """Yield non-header lines from notice body text."""
    for line in notice_text.strip().split("\n"):
        s = line.strip()
        if not s:
            yield ""
            continue
        if any(s.lower().startswith(p) for p in _SKIP_PREFIXES):
            continue
        yield s


# ── DOCX ─────────────────────────────────────────────────────────────────

def build_notice_docx(notice_text: str, taxpayer_name: str, pan: str, ay: str,
                      assessing_officer: str = "The Assessing Officer",
                      ward: str = "Ward / Circle",
                      ref_no: str = "",
                      din: str = "",
                      dsc: dict = None) -> BytesIO:
    """
    Build a formatted DOCX notice.
    Returns BytesIO ready for st.download_button.
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    doc  = Document()
    today = date.today().strftime("%d/%m/%Y")
    ref   = ref_no or f"ITO/{ay.replace('-','_')}/{pan}/142(1)"

    # Page margins
    for sec in doc.sections:
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.25)
        sec.right_margin  = Inches(1.25)

    # Letterhead
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("INCOME TAX DEPARTMENT")
    r.bold = True; r.font.size = Pt(14)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run(f"Government of India  —  {ward}").bold = True
    doc.add_paragraph()

    # DIN
    if din:
        doc.add_paragraph().add_run(f"DIN: {din}").bold = True

    # F.No and Date
    doc.add_paragraph().add_run(f"F.No.: {ref}").bold = True
    doc.add_paragraph().add_run(f"Date: {today}").bold = True
    doc.add_paragraph()

    # Subject
    doc.add_paragraph().add_run(
        f"Subject: Notice under Section 142(1) of the Income Tax Act, 1961 — "
        f"Assessment Year {ay} — in the case of {taxpayer_name} (PAN: {pan})."
    ).bold = True
    doc.add_paragraph()

    # Salutation + body
    doc.add_paragraph("Sir/Madam,")
    doc.add_paragraph()
    for line in _body_lines(notice_text):
        if not line:
            doc.add_paragraph()
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(3)

    doc.add_paragraph()
    doc.add_paragraph().add_run(
        "You are requested to furnish the above information/documents within "
        "15 days from the date of this notice. Non-compliance may attract "
        "penalty under Section 272A of the Income Tax Act, 1961."
    ).italic = True

    doc.add_paragraph(); doc.add_paragraph()
    doc.add_paragraph("Yours faithfully,")
    doc.add_paragraph(); doc.add_paragraph()
    doc.add_paragraph().add_run(assessing_officer).bold = True
    doc.add_paragraph(ward)
    doc.add_paragraph("Income Tax Department")
    doc.add_paragraph(f"Date: {today}")

    # DSC block
    if dsc:
        doc.add_paragraph()
        doc.add_paragraph("─" * 55)
        doc.add_paragraph().add_run("DIGITALLY SIGNED").bold = True
        for label, key in [
            ("Signed by",   "signed_by"),
            ("Designation", "designation"),
            ("Ward/Circle", "ward"),
            ("Date & Time", "signed_on"),
            ("Cert Serial", "cert_serial"),
            ("Valid Until", "valid_until"),
        ]:
            doc.add_paragraph(f"{label}: {dsc.get(key, '')}")
        doc.add_paragraph(dsc.get("note", ""))
        doc.add_paragraph("─" * 55)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── PDF ───────────────────────────────────────────────────────────────────

def build_notice_pdf(notice_text: str, taxpayer_name: str, pan: str, ay: str,
                     assessing_officer: str = "The Assessing Officer",
                     ward: str = "Ward / Circle",
                     ref_no: str = "",
                     din: str = "",
                     dsc: dict = None) -> BytesIO:
    """
    Build a formatted PDF notice using reportlab.
    Includes DIN header and DSC stamp box when provided.
    Returns BytesIO ready for st.download_button.
    """
    if not PDF_AVAILABLE:
        raise ImportError("reportlab not installed. Run: pip install reportlab")

    buf   = BytesIO()
    today = date.today().strftime("%d/%m/%Y")
    ref   = ref_no or f"ITO/{ay.replace('-','_')}/{pan}/142(1)"

    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=25*mm, rightMargin=20*mm,
        topMargin=22*mm, bottomMargin=20*mm,
        title=f"Notice 142(1) — {pan} — AY {ay}",
        author="Income Tax Department",
    )

    styles = getSampleStyleSheet()

    # Custom styles
    s_center = ParagraphStyle("center", parent=styles["Normal"],
                               alignment=TA_CENTER, fontSize=9)
    s_bold_c = ParagraphStyle("boldcenter", parent=styles["Normal"],
                               alignment=TA_CENTER, fontSize=11,
                               fontName="Helvetica-Bold")
    s_bold   = ParagraphStyle("bold", parent=styles["Normal"],
                               fontName="Helvetica-Bold", fontSize=9)
    s_normal = ParagraphStyle("norm", parent=styles["Normal"],
                               fontSize=9, leading=14, alignment=TA_JUSTIFY)
    s_italic = ParagraphStyle("ital", parent=styles["Normal"],
                               fontSize=8, fontName="Helvetica-Oblique",
                               leading=13, alignment=TA_JUSTIFY)
    s_small  = ParagraphStyle("small", parent=styles["Normal"],
                               fontSize=7.5, leading=11)
    s_din    = ParagraphStyle("din", parent=styles["Normal"],
                               fontSize=8, fontName="Helvetica-Bold",
                               alignment=TA_CENTER, textColor=colors.HexColor("#1a3a6b"))

    story = []

    # Letterhead
    story.append(Paragraph("INCOME TAX DEPARTMENT", s_bold_c))
    story.append(Spacer(1, 2*mm))
    story.append(Paragraph(f"Government of India  —  {ward}", s_center))
    story.append(Spacer(1, 1*mm))
    story.append(HRFlowable(width="100%", thickness=1.2,
                             color=colors.HexColor("#c0392b")))
    story.append(Spacer(1, 3*mm))

    # DIN
    if din:
        story.append(Paragraph(f"DIN: {din}", s_din))
        story.append(Spacer(1, 2*mm))

    # F.No and Date
    story.append(Paragraph(f"<b>F.No.:</b> {ref}", s_normal))
    story.append(Paragraph(f"<b>Date:</b> {today}", s_normal))
    story.append(Spacer(1, 4*mm))

    # Subject
    story.append(Paragraph(
        f"<b>Subject:</b> Notice under Section 142(1) of the Income Tax Act, 1961 — "
        f"Assessment Year {ay} — in the case of {taxpayer_name} (PAN: {pan}).",
        s_bold
    ))
    story.append(Spacer(1, 4*mm))

    # Salutation
    story.append(Paragraph("Sir/Madam,", s_normal))
    story.append(Spacer(1, 3*mm))

    # Body
    for line in _body_lines(notice_text):
        if not line:
            story.append(Spacer(1, 3*mm))
        else:
            story.append(Paragraph(line, s_normal))

    story.append(Spacer(1, 4*mm))

    # Compliance deadline
    story.append(Paragraph(
        "You are requested to furnish the above information/documents within "
        "15 days from the date of this notice. Non-compliance may attract "
        "penalty under Section 272A of the Income Tax Act, 1961.",
        s_italic
    ))
    story.append(Spacer(1, 8*mm))

    # Signature block
    story.append(Paragraph("Yours faithfully,", s_normal))
    story.append(Spacer(1, 10*mm))
    story.append(Paragraph(f"<b>{assessing_officer}</b>", s_normal))
    story.append(Paragraph(ward, s_normal))
    story.append(Paragraph("Income Tax Department", s_normal))
    story.append(Paragraph(f"Date: {today}", s_normal))

    # DSC stamp box
    if dsc:
        story.append(Spacer(1, 8*mm))
        dsc_data = [
            ["DIGITALLY SIGNED", ""],
            ["Signed by",   dsc.get("signed_by", "")],
            ["Designation", dsc.get("designation", "")],
            ["Ward/Circle", dsc.get("ward", "")],
            ["Date & Time", dsc.get("signed_on", "")],
            ["Cert. Serial",dsc.get("cert_serial", "")],
            ["Valid Until", dsc.get("valid_until", "")],
            [dsc.get("note", ""), ""],
        ]
        tbl = Table(dsc_data, colWidths=[38*mm, 120*mm])
        tbl.setStyle(TableStyle([
            ("BACKGROUND",   (0, 0), (-1, 0), colors.HexColor("#d6eaf8")),
            ("FONTNAME",     (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE",     (0, 0), (-1, 0), 8),
            ("SPAN",         (0, 0), (-1, 0)),
            ("ALIGN",        (0, 0), (-1, 0), "CENTER"),
            ("FONTSIZE",     (0, 1), (-1, -1), 7.5),
            ("FONTNAME",     (0, 1), (0, -1), "Helvetica-Bold"),
            ("TEXTCOLOR",    (0, 1), (0, -1), colors.HexColor("#1a3a6b")),
            ("BACKGROUND",   (0, 1), (-1, -1), colors.HexColor("#f0f8ff")),
            ("GRID",         (0, 0), (-1, -1), 0.4, colors.HexColor("#2980b9")),
            ("ROWBACKGROUND",(0, -1), (-1, -1), colors.HexColor("#e8f4f8")),
            ("FONTNAME",     (0, -1), (-1, -1), "Helvetica-Oblique"),
            ("FONTSIZE",     (0, -1), (-1, -1), 7),
            ("SPAN",         (0, -1), (-1, -1)),
            ("TOPPADDING",   (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING",(0, 0), (-1, -1), 3),
            ("LEFTPADDING",  (0, 0), (-1, -1), 5),
        ]))
        story.append(tbl)

    doc.build(story)
    buf.seek(0)
    return buf
